"""Import IT/computer repair catalog from docx into the database.

Parses 'Доки - свежий прайс (1).docx' and creates:
  - Profession: IT (Компьютерный мастер)
  - Groups & Subgroups based on service categories
  - ServiceItems for each row

Usage:
  python -m scripts.import_it_catalog
  python -m scripts.import_it_catalog --dry-run
"""

import argparse
import asyncio
import logging
import re
import sys
from pathlib import Path

try:
    import docx
except ImportError:
    print("python-docx не установлен: pip install python-docx")
    sys.exit(1)

if sys.platform == "win32":
    sys.stdout.reconfigure(encoding="utf-8")

logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)-8s %(message)s", datefmt="%H:%M:%S")
logger = logging.getLogger("import_it")

DOCX_PATH = Path(__file__).resolve().parent.parent / "Доки - свежий прайс (1).docx"

# --- Categorization of services by row ranges ---
# Manually mapped from the docx structure for quality grouping

GROUPS = {
    "service_diagnostika": {
        "name": "Сервис и диагностика",
        "subgroups": {
            "vyezd": {
                "name": "Выезд и диагностика",
                "rows": [1, 2, 3, 4, 5, 6, 7, 8, 9],
            },
        },
    },
    "remont_pk": {
        "name": "Ремонт ПК и ноутбуков",
        "subgroups": {
            "obsluzhivanie": {
                "name": "Обслуживание и чистка",
                "rows": [10, 11, 12, 13, 14, 15, 19, 21, 22],
            },
            "remont_plat": {
                "name": "Ремонт плат и компонентов",
                "rows": [16, 17, 18],
            },
            "komplektuyushchie": {
                "name": "Комплектующие и подбор",
                "rows": [20, 23],
            },
        },
    },
    "windows_os": {
        "name": "Windows и ОС",
        "subgroups": {
            "ustanovka_windows": {
                "name": "Установка и восстановление Windows",
                "rows": [24, 25, 26, 27, 29, 33, 36],
            },
            "nastroyka_os": {
                "name": "Настройка и оптимизация ОС",
                "rows": [28, 30, 31, 32, 34, 35],
            },
        },
    },
    "drajvery_po": {
        "name": "Драйверы и ПО",
        "subgroups": {
            "drajvery": {
                "name": "Драйверы",
                "rows": [37, 38, 39],
            },
            "ustanovka_po": {
                "name": "Установка программ",
                "rows": [40, 41, 42, 43, 44, 45, 46, 47, 48],
            },
            "pakety_po": {
                "name": "Пакеты ПО",
                "rows": [49, 50, 51, 52, 53, 54],
            },
        },
    },
    "hdd_ssd": {
        "name": "Жёсткие диски и данные",
        "subgroups": {
            "rabota_hdd": {
                "name": "Работа с HDD/SSD",
                "rows": [55, 57, 58, 62, 63],
            },
            "vosstanovlenie_dannyh": {
                "name": "Восстановление данных",
                "rows": [56, 59, 60, 61],
            },
            "bios": {
                "name": "BIOS и прошивки",
                "rows": [64, 65, 66],
            },
        },
    },
    "bezopasnost": {
        "name": "Безопасность и антивирусы",
        "subgroups": {
            "virusy": {
                "name": "Удаление вирусов",
                "rows": [69, 70, 71, 72, 73, 74, 82],
            },
            "antivirus": {
                "name": "Антивирусное ПО",
                "rows": [75, 76, 77, 78, 79, 80, 81],
            },
            "zashchita": {
                "name": "Защита и шифрование",
                "rows": [67, 68],
            },
        },
    },
    "seti_internet": {
        "name": "Сети и интернет",
        "subgroups": {
            "nastroyka_inet": {
                "name": "Настройка интернета",
                "rows": [83, 84, 85, 86, 87, 88],
            },
            "wifi_routery": {
                "name": "Wi-Fi и роутеры",
                "rows": [89, 90, 92, 93, 94, 95, 96, 97, 98, 99, 100, 101, 102, 103, 104],
            },
            "registratsiya": {
                "name": "Регистрация и аккаунты",
                "rows": [91],
            },
        },
    },
    "planshety_mobilnye": {
        "name": "Планшеты и мобильные",
        "subgroups": {
            "android_wp": {
                "name": "Android / Windows Phone",
                "rows": [105, 106, 107, 108],
            },
            "ip_tv": {
                "name": "IP-TV и приставки",
                "rows": [109],
            },
            "diagnostika_apparatnaya": {
                "name": "Аппаратная диагностика",
                "rows": [110],
            },
        },
    },
    "apple": {
        "name": "Apple и macOS",
        "subgroups": {
            "uchetnaya_zapis": {
                "name": "Учётные записи и iCloud",
                "rows": [111, 123],
            },
            "macos_ustanovka": {
                "name": "Установка и настройка macOS",
                "rows": [112, 113, 114, 115, 116, 117, 118, 119],
            },
            "apple_servisy": {
                "name": "Сервисы Apple",
                "rows": [120, 121, 122],
            },
        },
    },
    "konsoli": {
        "name": "Игровые консоли",
        "subgroups": {
            "remont_konsoli": {
                "name": "Ремонт и обслуживание консолей",
                "rows": [125, 126, 127, 128, 129, 130, 131, 132],
            },
        },
    },
    "dostavka": {
        "name": "Доставка и логистика",
        "subgroups": {
            "dostavka_oborud": {
                "name": "Доставка оборудования",
                "rows": [124],
            },
        },
    },
    "macbook_remont": {
        "name": "Ремонт MacBook",
        "subgroups": {
            "macbook_pro": {
                "name": "MacBook Pro",
                "tables": [1, 2],
            },
            "macbook_air": {
                "name": "MacBook Air",
                "tables": [3, 4],
            },
            "macbook_retina": {
                "name": "MacBook Retina",
                "tables": [5, 6],
            },
        },
    },
}

# Additional hashtags per group for better search
GROUP_TAGS = {
    "service_diagnostika": "#компьютерныймастер #диагностика #выезд",
    "remont_pk": "#компьютерныймастер #ремонтпк #ноутбук #чистка",
    "windows_os": "#компьютерныймастер #windows #ос #установка",
    "drajvery_po": "#компьютерныймастер #драйверы #программы #софт",
    "hdd_ssd": "#компьютерныймастер #жесткийдиск #ssd #hdd #данные",
    "bezopasnost": "#компьютерныймастер #вирусы #антивирус #безопасность",
    "seti_internet": "#компьютерныймастер #интернет #wifi #роутер #сеть",
    "planshety_mobilnye": "#компьютерныймастер #планшет #android #мобильный",
    "apple": "#компьютерныймастер #apple #mac #macos #icloud",
    "konsoli": "#компьютерныймастер #консоль #playstation #xbox",
    "dostavka": "#компьютерныймастер #доставка",
    "macbook_remont": "#компьютерныймастер #macbook #apple #ремонт",
}


def parse_price(text: str) -> int:
    """Extract numeric price from string like '1 490 руб.' or '1490руб.'"""
    text = text.strip().lower().replace("руб.", "").replace("₽", "").replace("\xa0", "")
    text = re.sub(r"\s+", "", text)
    text = text.replace(",", ".")
    if not text or text == "0" or "уточн" in text or "недоступ" in text:
        return 0
    # Handle "от 4900"
    text = text.replace("от", "").strip()
    # Handle "500р+500р/клавиша" -> take first number
    if "+" in text:
        text = text.split("+")[0]
    text = re.sub(r"р$", "", text)
    try:
        return int(float(text))
    except (ValueError, TypeError):
        return 0


def make_code(group_code: str, idx: int) -> str:
    """Generate a short unique item code."""
    prefix = group_code.upper().replace("_", "-")
    return f"IT-{prefix}-{idx:03d}"


def build_search_text(name: str, description: str, hashtags: str) -> str:
    parts = [
        name.lower(),
        description.lower() if description else "",
        hashtags.replace("#", " ").lower() if hashtags else "",
    ]
    return " ".join(filter(None, parts)).strip()


def parse_main_table(doc) -> dict[int, dict]:
    """Parse Table 0 into {row_num: {name, price, description, unit}}"""
    table = doc.tables[0]
    items = {}
    for ri, row in enumerate(table.rows):
        if ri == 0:
            continue
        cells = [c.text.strip().replace("\n", " ") for c in row.cells]
        try:
            row_num = int(cells[0])
        except (ValueError, IndexError):
            continue
        name = cells[1].strip() if len(cells) > 1 else ""
        price_text = cells[2].strip() if len(cells) > 2 else "0"
        description = cells[3].strip() if len(cells) > 3 else ""
        price = parse_price(price_text)

        items[row_num] = {
            "name": name,
            "price": price,
            "description": description,
            "unit": "усл.",
        }
    return items


def parse_macbook_tables(doc) -> list[dict]:
    """Parse MacBook tables (1-6) into flat service items."""
    items = []
    # Tables 1,3,5 have: Series, Diagonal, Model, Marking, then repair types with prices
    # Tables 2,4,6 have additional repair types

    for table_idx in [1, 2, 3, 4, 5, 6]:
        if table_idx >= len(doc.tables):
            continue
        table = doc.tables[table_idx]
        header_cells = [c.text.strip().replace("\n", " ") for c in table.rows[0].cells]

        # Find service columns (skip first 4: series, diagonal, model, marking)
        service_cols = header_cells[4:]

        for ri in range(1, len(table.rows)):
            row_cells = [c.text.strip().replace("\n", " ") for c in table.rows[ri].cells]
            series = row_cells[0] if len(row_cells) > 0 else ""
            diagonal = row_cells[1] if len(row_cells) > 1 else ""
            model_years = row_cells[2] if len(row_cells) > 2 else ""
            marking = row_cells[3] if len(row_cells) > 3 else ""

            for ci, service_name in enumerate(service_cols):
                price_text = row_cells[4 + ci] if (4 + ci) < len(row_cells) else ""
                price = parse_price(price_text)
                if price <= 0:
                    continue

                item_name = f"{service_name} — {series} {diagonal}\" ({marking})"
                items.append({
                    "name": item_name,
                    "price": price,
                    "description": f"{model_years}. {series} {diagonal}\"",
                    "unit": "усл.",
                    "series": series,
                    "marking": marking,
                    "table_idx": table_idx,
                })

    return items


async def run_import(dry_run: bool = False):
    from sqlalchemy import delete, select, func
    from app.database import get_async_session
    from app.models.catalog import (
        Profession, ServiceGroup, ServiceItem, ServiceSubgroup,
    )

    docx_path = DOCX_PATH
    if not docx_path.exists():
        # Try current directory
        docx_path = Path("Доки - свежий прайс (1).docx")
    if not docx_path.exists():
        docx_path = Path("catalog_it.docx")
    if not docx_path.exists():
        logger.error("Docx file not found: %s", DOCX_PATH)
        sys.exit(1)

    logger.info("Открываем %s...", docx_path.name)
    doc = docx.Document(str(docx_path))
    logger.info("Таблиц: %d, параграфов: %d", len(doc.tables), len(doc.paragraphs))

    # Parse data
    main_items = parse_main_table(doc)
    macbook_items = parse_macbook_tables(doc)
    logger.info("Основные услуги: %d строк", len(main_items))
    logger.info("MacBook услуги: %d позиций", len(macbook_items))

    if dry_run:
        logger.info("DRY RUN — примеры:")
        for row_num in [1, 10, 50, 100, 132]:
            if row_num in main_items:
                item = main_items[row_num]
                logger.info("  #%d: %s — %d₽", row_num, item["name"], item["price"])
        logger.info("MacBook примеры:")
        for mb in macbook_items[:5]:
            logger.info("  %s — %d₽", mb["name"], mb["price"])
        return

    session_factory = get_async_session()
    async with session_factory() as session:
        # Check if IT profession already exists
        existing_prof = (await session.execute(
            select(Profession).where(Profession.code == "IT")
        )).scalar_one_or_none()

        if existing_prof:
            logger.info("Профессия IT уже существует (id=%d), удаляем старые данные...", existing_prof.id)
            await session.execute(
                delete(ServiceItem).where(ServiceItem.profession_id == existing_prof.id)
            )
            await session.execute(
                delete(ServiceSubgroup).where(
                    ServiceSubgroup.group_id.in_(
                        select(ServiceGroup.id).where(ServiceGroup.profession_id == existing_prof.id)
                    )
                )
            )
            await session.execute(
                delete(ServiceGroup).where(ServiceGroup.profession_id == existing_prof.id)
            )
            await session.commit()
            prof = existing_prof
        else:
            prof = Profession(
                code="IT",
                name="Компьютерный мастер",
                icon="💻",
                sort_priority=3,
            )
            session.add(prof)
            await session.flush()

        logger.info("Профессия: IT — %s (id=%d)", prof.name, prof.id)

        # Create groups and subgroups
        groups_map = {}  # group_code -> group_id
        subgroups_map = {}  # (group_code, sub_code) -> sub_id
        sort_g = 0
        sort_s = 0

        for group_code, group_data in GROUPS.items():
            grp = ServiceGroup(
                profession_id=prof.id,
                code=f"it_{group_code}",
                name=group_data["name"],
                sort_priority=sort_g,
            )
            session.add(grp)
            await session.flush()
            groups_map[group_code] = grp.id
            sort_g += 1

            for sub_code, sub_data in group_data["subgroups"].items():
                sub = ServiceSubgroup(
                    group_id=grp.id,
                    code=f"it_{group_code}_{sub_code}",
                    name=sub_data["name"],
                    sort_priority=sort_s,
                )
                session.add(sub)
                await session.flush()
                subgroups_map[(group_code, sub_code)] = sub.id
                sort_s += 1

        logger.info("Групп: %d, подгрупп: %d", len(groups_map), len(subgroups_map))

        # Build reverse map: row_num -> (group_code, sub_code)
        row_to_location = {}
        for group_code, group_data in GROUPS.items():
            for sub_code, sub_data in group_data["subgroups"].items():
                for row_num in sub_data.get("rows", []):
                    row_to_location[row_num] = (group_code, sub_code)

        # Insert main table items
        items_added = 0
        sort_order = 0

        for row_num, item_data in sorted(main_items.items()):
            loc = row_to_location.get(row_num)
            if not loc:
                logger.warning("  Строка %d не привязана к группе: %s", row_num, item_data["name"][:50])
                continue

            group_code, sub_code = loc
            grp_id = groups_map[group_code]
            sub_id = subgroups_map[(group_code, sub_code)]
            tags = GROUP_TAGS.get(group_code, "#компьютерныймастер")

            code = f"IT-{row_num:03d}"
            slug = re.sub(r"[^a-z0-9]+", "_", item_data["name"].lower()[:60].strip())[:60]

            price = item_data["price"]
            # For services with 0 price (free callout, free diag), set price range
            price_min = price
            price_max = price
            price_rec = price

            search_text = build_search_text(item_data["name"], item_data["description"], tags)

            si = ServiceItem(
                sort_order=sort_order,
                profession_id=prof.id,
                group_id=grp_id,
                subgroup_id=sub_id,
                code=code,
                slug=slug,
                name=item_data["name"],
                description=item_data["description"][:500] if item_data["description"] else None,
                unit=item_data["unit"],
                price_min=price_min,
                price_max=price_max,
                price_recommended=price_rec,
                currency="RUB",
                record_type="atomic",
                calc_strategy="PER_UNIT",
                selection_mode="single",
                complexity="std",
                confidence="HIGH",
                labor_only=True,
                hashtags=tags,
                search_text=search_text,
                city="Стерлитамак",
                region="Башкортостан",
                note=item_data["description"][:200] if item_data["description"] else None,
                is_active=True,
            )
            session.add(si)
            items_added += 1
            sort_order += 1

            if items_added % 50 == 0:
                await session.flush()
                logger.info("  ... основных: %d", items_added)

        await session.flush()
        logger.info("Основных позиций добавлено: %d", items_added)

        # Insert MacBook items
        mb_added = 0
        # Deduplicate by (service_name, marking) — use average price if multiple
        mb_seen = set()

        for mb in macbook_items:
            table_idx = mb["table_idx"]
            series = mb.get("series", "")

            # Determine subgroup
            if table_idx in [1, 2]:
                sub_key = ("macbook_remont", "macbook_pro")
            elif table_idx in [3, 4]:
                sub_key = ("macbook_remont", "macbook_air")
            else:
                sub_key = ("macbook_remont", "macbook_retina")

            grp_id = groups_map["macbook_remont"]
            sub_id = subgroups_map[sub_key]
            tags = GROUP_TAGS.get("macbook_remont", "#macbook")

            # Create unique code
            code = f"IT-MB-{table_idx:d}-{mb_added:03d}"

            # Skip exact duplicates
            dedup_key = (mb["name"], mb["price"])
            if dedup_key in mb_seen:
                continue
            mb_seen.add(dedup_key)

            slug = re.sub(r"[^a-z0-9]+", "_", mb["name"].lower()[:60])[:60]
            search_text = build_search_text(mb["name"], mb["description"], tags)

            si = ServiceItem(
                sort_order=sort_order,
                profession_id=prof.id,
                group_id=grp_id,
                subgroup_id=sub_id,
                code=code,
                slug=slug,
                name=mb["name"],
                description=mb["description"],
                unit="усл.",
                price_min=mb["price"],
                price_max=mb["price"],
                price_recommended=mb["price"],
                currency="RUB",
                record_type="atomic",
                calc_strategy="PER_UNIT",
                selection_mode="single",
                complexity="complex",
                confidence="HIGH",
                labor_only=True,
                hashtags=tags,
                search_text=search_text,
                city="Стерлитамак",
                region="Башкортостан",
                is_active=True,
            )
            session.add(si)
            mb_added += 1
            sort_order += 1

            if mb_added % 50 == 0:
                await session.flush()

        await session.flush()
        logger.info("MacBook позиций добавлено: %d", mb_added)

        await session.commit()

        # Summary
        total = (await session.execute(
            select(func.count(ServiceItem.id)).where(ServiceItem.profession_id == prof.id)
        )).scalar()

        logger.info("=== Импорт IT-каталога завершён ===")
        logger.info("  Профессия: %s (id=%d)", prof.name, prof.id)
        logger.info("  Групп: %d", len(groups_map))
        logger.info("  Подгрупп: %d", len(subgroups_map))
        logger.info("  Позиций всего: %d (основных: %d, MacBook: %d)", total, items_added, mb_added)


def main():
    parser = argparse.ArgumentParser(description="Импорт IT-каталога из docx")
    parser.add_argument("--dry-run", action="store_true")
    args = parser.parse_args()
    asyncio.run(run_import(dry_run=args.dry_run))


if __name__ == "__main__":
    main()
